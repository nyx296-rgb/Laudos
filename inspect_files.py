
import openpyxl
from docx import Document
import re

# Inspect DOCX
print('=== DOCX PARAGRAPHS ===')
doc = Document(r'c:\Users\Eu\Documents\Laudos\Modelaudo_FORMS.docx')
for i, para in enumerate(doc.paragraphs):
    print(f'P{i}: {repr(para.text)}')

print('\n=== DOCX TABLES ===')
for ti, table in enumerate(doc.tables):
    print(f'--- Table {ti} ---')
    for ri, row in enumerate(table.rows):
        row_data = []
        for ci, cell in enumerate(row.cells):
            row_data.append(repr(cell.text.strip()[:100]))
        print(f'  Row {ri}: {row_data}')

# Extract all tags
all_text = []
for para in doc.paragraphs:
    all_text.append(para.text)
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                all_text.append(para.text)
full_text = '\n'.join(all_text)
tags = re.findall(r'\{\{[^}]+\}\}', full_text)
unique_tags = sorted(set(tags))
print('\n=== UNIQUE TAGS ===')
for tag in unique_tags:
    print(tag)

# Inspect XLSX
print('\n=== XLSX DATA ===')
wb = openpyxl.load_workbook(r'c:\Users\Eu\Documents\Laudos\cod_tasy.xlsx')
for sheet in wb.sheetnames:
    ws = wb[sheet]
    print(f'Sheet: {sheet}, rows: {ws.max_row}, cols: {ws.max_column}')
    for i, row in enumerate(ws.iter_rows(values_only=True)):
        if any(v is not None for v in row):
            print(f'  Row {i}: {list(row)}')
