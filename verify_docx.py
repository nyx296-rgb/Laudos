from docx import Document

def verify_docx(path):
    doc = Document(path)
    print(f"Verificando: {path}")
    
    # Check paragraphs
    print("--- Texto nos parágrafos ---")
    for p in doc.paragraphs:
        if 'TEST_999' in p.text or 'Antigravity' in p.text:
            print(f"P: {p.text}")
            
    # Check tables
    print("--- Texto nas tabelas ---")
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells]
            if any('Notebook' in t or 'Monitor' in t for t in row_text):
                print(f"Row: {row_text}")

if __name__ == '__main__':
    verify_docx('laudo_TEST_999.docx')
