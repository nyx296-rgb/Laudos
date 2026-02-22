"""
Laudo Generator - Core logic for filling DOCX template and converting to PDF.
"""
import os
import copy
import tempfile
import qrcode
from io import BytesIO
from docx import Document
from docx.shared import Pt, Cm
from lxml import etree

# Path to the template
TEMPLATE_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Modelaudo_FORMS.docx')


# ─────────────────────────────────────────────
# LOW-LEVEL HELPERS
# ─────────────────────────────────────────────

def _xml_replace(xml_element, tag_map):
    """Replace all tag occurrences in the XML of an element in-place."""
    xml_str = etree.tostring(xml_element, encoding='unicode')
    for tag, value in tag_map.items():
        xml_str = xml_str.replace(tag, str(value) if value is not None else '')
    return etree.fromstring(xml_str)


def replace_in_paragraph(paragraph, replacements):
    """Replace {{tags}} in a single paragraph, handling runs split across multiple pieces."""
    full_text = ''.join(run.text for run in paragraph.runs)
    new_text = full_text
    for tag, value in replacements.items():
        new_text = new_text.replace(tag, str(value) if value is not None else '')

    if new_text != full_text and paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ''


def replace_in_doc(doc, replacements):
    """Replace tags in every paragraph, table cell, header, and footer of the document."""
    # Main document body
    for para in doc.paragraphs:
        replace_in_paragraph(para, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_paragraph(para, replacements)

    # Headers and Footers
    for section in doc.sections:
        # Header
        for para in section.header.paragraphs:
            replace_in_paragraph(para, replacements)
        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        replace_in_paragraph(para, replacements)
        
        # Footer
        for para in section.footer.paragraphs:
            replace_in_paragraph(para, replacements)
        for table in section.footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        replace_in_paragraph(para, replacements)


# ─────────────────────────────────────────────
# EQUIPMENT TABLE FILLER
# ─────────────────────────────────────────────

def fill_equipment_table(doc, equipamentos):
    """
    The template has one master table (6 rows × 10 cols).
    Row 5 contains the {{Tasy}}, {{Item}}, {{Marca}}, {{Modelo}},
    {{Nº. Série}}, {{Quantidade}}, {{Situação do Equipamento}} tags.

    Strategy:
    - For each equipment item, clone the template row and replace tags with real data.
    - Remove the original template row at the end.
    - If no equipment provided, clear the tags from the template row.
    """
    equip_table = None
    template_row_index = None

def fill_equipment_table(doc, equipamentos):
    """
    Consolidates text nodes in the template row to prevent fragmented tags,
    duplicates the row for each item, and fills the tags.
    """
    from docx.table import _Row
    
    template_row = None
    target_table = None
    
    # 1. Search for the row containing equipment row tags (prioritize Tasy)
    for table in doc.tables:
        for row in table.rows:
            row_text = "".join(cell.text for cell in row.cells)
            if "{{Tasy}}" in row_text or "{{Quantidade}}" in row_text:
                target_table = table
                template_row = row
                break
        if template_row:
            break

    if not template_row:
        return

    # 2. Replicate and fill
    template_tr_xml = template_row._tr
    parent_xml = template_tr_xml.getparent()

    for equip in equipamentos:
        # Clone the template XML
        new_tr_xml = copy.deepcopy(template_tr_xml)
        
        # Consolidation: Merge fragmented <w:t> nodes within each <w:p> in the row
        # to ensure tags like {{Tasy}} are not split.
        _consolidate_xml_text(new_tr_xml)
        
        # Wrap in Row object to use our paragraph replacer
        # (which is already logic-heavy for run handling)
        new_row = _Row(new_tr_xml, target_table)
        tag_map = _build_equip_replacements(equip)
        for cell in new_row.cells:
            for para in cell.paragraphs:
                replace_in_paragraph(para, tag_map)
        
        # Insert into the document XML
        template_tr_xml.addprevious(new_tr_xml)

    # 3. Always remove the placeholder row
    parent_xml.remove(template_tr_xml)


def _consolidate_xml_text(element):
    """
    Very basic XML consolidation: in each paragraph (<w:p>), 
    it tries to merge text into a single run if tags are fragmented.
    For our purpose, simply ensuring each <w:p> text is handled works.
    """
    from docx.oxml.ns import qn
    for p in element.findall('.//' + qn('w:p')):
        # Get all runs in this paragraph
        runs = p.findall('.//' + qn('w:r'))
        if not runs:
            continue
        
        # Collect all text
        full_text = ""
        for r in runs:
            for t in r.findall('.//' + qn('w:t')):
                if t.text: full_text += t.text
        
        # If we see a tag split, we simplify
        if "{{" in full_text and "}}" in full_text:
            # Clear all runs
            for r in runs:
                p.remove(r)
            # Add back a single run with full text
            from docx.oxml.shared import OxmlElement
            new_r = OxmlElement('w:r')
            new_t = OxmlElement('w:t')
            new_t.text = full_text
            new_r.append(new_t)
            p.append(new_r)


def _build_equip_replacements(equip):
    return {
        # Per user request: map item name to {{Tasy}} and also {{Item}} for local rows
        '{{Tasy}}':                       equip.get('item', ''),
        '{{Item}}':                        equip.get('item', ''),
        '{{item}}':                        equip.get('item', ''),
        '{{Marca}}':                       equip.get('marca', ''),
        '{{Modelo}}':                      equip.get('modelo', ''),
        '{{Nº. Série}}':                   equip.get('serie', ''),
        '{{Nº Série}}':                    equip.get('serie', ''),
        '{{Quantidade}}':                  str(equip.get('quantidade', '1')),
        '{{Qtd}}':                         str(equip.get('quantidade', '1')),
        '{{Situação do Equipamento}}':     equip.get('situacao', ''),
        '{{Situação do Equip.}}':          equip.get('situacao', ''),
    }


# ─────────────────────────────────────────────
# QR CODE
# ─────────────────────────────────────────────

def add_qrcode_to_doc(doc, laudo_id, verification_url=None):
    """Generate a QR Code image and append it to the document."""
    if verification_url is None:
        verification_url = (
            f"LAUDO TÉCNICO Nº {laudo_id} — "
            "Documento gerado eletronicamente. Verifique a autenticidade no setor responsável."
        )

    qr = qrcode.QRCode(
        version=1,
        error_correction=qrcode.constants.ERROR_CORRECT_M,
        box_size=10,
        border=4,
    )
    qr.add_data(verification_url)
    qr.make(fit=True)
    img = qr.make_image(fill_color='black', back_color='white')

    img_bytes = BytesIO()
    img.save(img_bytes, format='PNG')
    img_bytes.seek(0)

    qr_para = doc.add_paragraph()
    qr_para.alignment = 2  # WD_ALIGN_PARAGRAPH.RIGHT
    run = qr_para.add_run()
    run.add_picture(img_bytes, width=Cm(3.5))

    caption = doc.add_paragraph('QR Code de Autenticidade')
    caption.alignment = 2
    for run in caption.runs:
        run.font.size = Pt(8)


# ─────────────────────────────────────────────
# PDF CONVERSION
# ─────────────────────────────────────────────

def _try_convert_pdf(docx_path, pdf_path):
    """
    Try multiple methods to convert DOCX → PDF.
    Returns True on success, False if not possible.
    """
    # 1. Try docx2pdf (uses Microsoft Word on Windows)
    try:
        from docx2pdf import convert
        convert(docx_path, pdf_path)
        if os.path.exists(pdf_path):
            return True
    except Exception as e:
        print(f"[docx2pdf] failed: {e}")

    # 2. Try LibreOffice
    import subprocess, shutil as sh
    soffice = sh.which('soffice') or sh.which('libreoffice')
    if not soffice:
        for candidate in [
            r'C:\Program Files\LibreOffice\program\soffice.exe',
            r'C:\Program Files (x86)\LibreOffice\program\soffice.exe',
        ]:
            if os.path.exists(candidate):
                soffice = candidate
                break

    if soffice:
        try:
            out_dir = os.path.dirname(pdf_path)
            result = subprocess.run(
                [soffice, '--headless', '--convert-to', 'pdf', '--outdir', out_dir, docx_path],
                capture_output=True, timeout=60
            )
            # LibreOffice names the output file <docx_basename>.pdf
            lo_pdf = os.path.join(out_dir, os.path.splitext(os.path.basename(docx_path))[0] + '.pdf')
            if os.path.exists(lo_pdf):
                if lo_pdf != pdf_path:
                    os.rename(lo_pdf, pdf_path)
                return True
        except Exception as e:
            print(f"[LibreOffice] failed: {e}")

    return False


# ─────────────────────────────────────────────
# MAIN ENTRY POINT
# ─────────────────────────────────────────────

def generate_laudo(data):
    """
    Generate a laudo or purchase request document.

    Args:
        data (dict): Form fields from the frontend.

    Returns:
        (output_path, temp_dir): Path to the generated file (PDF or DOCX) 
        and the temp directory (caller should clean up).
    """
    tipo = data.get('tipo', 'laudo')
    template_filename = 'ModelSolicomp_FORMS.docx' if tipo == 'compra' else 'Modelaudo_FORMS.docx'
    current_template_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), template_filename)
    
    doc = Document(current_template_path)

    laudo_id = data.get('id_laudo', 'laudo')

    # Fill equipment table rows first (they need their own template row cloning)
    equipamentos = data.get('equipamentos', [])
    if equipamentos:
        fill_equipment_table(doc, equipamentos)

    # Replace scalar tags
    # Comprehensive mapping for all global placeholders
    replacements = {
        '{{ID_Laudo}}':         laudo_id,
        '{{Laudo}}':            laudo_id,
        '{{Data}}':             data.get('data', ''),
        '{{Local}}':            data.get('local', ''),
        '{{Unidade}}':          data.get('unidade', ''),
        '{{Setor}}':            data.get('setor', ''),
        '{{setor}}':            data.get('setor', ''),
        '{{Problema}}':         data.get('descricao_problema', ''),
        '{{Descrição do problema:}}': data.get('descricao_problema', ''),
        '{{Atendimento:}}':     data.get('descricao_problema', ''),
        '{{Marca}}':            data.get('marca', ''),
        '{{Modelo}}':           data.get('modelo', ''),
        '{{Nº Série}}':         data.get('serie', ''),
        '{{Nº. Série}}':        data.get('serie', ''),
        '{{Nº de Série}}':      data.get('serie', ''),
        '{{Situação do Equipamento}}': data.get('situacao', ''),
        '{{Situação do Equipemento}}': data.get('situacao', ''), # Fix user typo in template
        '{{Situação do Equip.}}': data.get('situacao', ''),
        '{{Situação}}':         data.get('situacao', ''),
        '{{Item}}':             data.get('item_defeito', ''),
        '{{item}}':             data.get('item_defeito', ''),
        '{{Tasy}}':             data.get('equipamentos', [{}])[0].get('item', '') if data.get('equipamentos') else '',
        '{{Quantidade}}':       str(data.get('equipamentos', [{}])[0].get('quantidade', '')) if data.get('equipamentos') else '',
        '{{Qtd}}':              str(data.get('equipamentos', [{}])[0].get('quantidade', '')) if data.get('equipamentos') else '',
        '{{Nome do Analista}}': data.get('nome_analista', ''),
        '{{Cargo do Analista}}':data.get('cargo_analista', ''),
        '{{Preenchido por}}':  data.get('preenchido_por', ''),
        '{{Preenchido}}':      data.get('preenchido_por', ''),
    }
    
    # Run replacement across the whole document
    replace_in_doc(doc, replacements)

    # Add QR Code
    verificacao = data.get(
        'verificacao_url',
        f"Laudo Técnico Nº {laudo_id}"
    )
    add_qrcode_to_doc(doc, laudo_id, verificacao)

    # Save to temp dir
    temp_dir = tempfile.mkdtemp()
    
    # Sanitize laudo_id for filename (replace / with -)
    safe_laudo_id = laudo_id.replace('/', '-')
    
    docx_path = os.path.join(temp_dir, f'laudo_{safe_laudo_id}.docx')
    pdf_path  = os.path.join(temp_dir, f'laudo_{safe_laudo_id}.pdf')
    doc.save(docx_path)

    # Try PDF conversion
    if _try_convert_pdf(docx_path, pdf_path):
        return pdf_path, temp_dir

    # Fallback: return DOCX
    print("[INFO] Returning DOCX (no PDF converter available).")
    return docx_path, temp_dir
